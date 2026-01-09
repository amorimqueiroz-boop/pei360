import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt
from openai import OpenAI
from pypdf import PdfReader
from fpdf import FPDF
import base64
import os
import re

# --- 1. CONFIGURAÇÃO INICIAL ---
def get_favicon():
    if os.path.exists("iconeaba.png"): return "iconeaba.png"
    return "📘"

st.set_page_config(
    page_title="PEI 360º",
    page_icon=get_favicon(),
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 2. UTILITÁRIOS ---
def finding_logo():
    possiveis = ["360.png", "360.jpg", "logo.png", "logo.jpg", "iconeaba.png"]
    for nome in possiveis:
        if os.path.exists(nome): return nome
    return None

def get_base64_image(image_path):
    if not image_path: return ""
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

def ler_pdf(arquivo):
    if arquivo is None: return ""
    try:
        reader = PdfReader(arquivo)
        texto = ""
        for i, page in enumerate(reader.pages):
            if i >= 6: break 
            texto += page.extract_text() + "\n"
        return texto
    except Exception as e: return f"Erro ao ler PDF: {e}"

def limpar_texto_pdf(texto):
    if not texto: return ""
    texto = texto.replace('**', '').replace('__', '')
    texto = texto.replace('### ', '').replace('## ', '').replace('# ', '')
    texto = texto.replace('* ', '-') 
    texto = texto.replace('–', '-').replace('—', '-')
    texto = texto.replace('“', '"').replace('”', '"').replace('‘', "'").replace('’', "'")
    texto = re.sub(r'[^\x00-\xff]', '', texto) 
    return texto

# --- 3. CSS (DESIGN SYSTEM PREMIUM) ---
st.markdown("""
    <link href="https://cdn.jsdelivr.net/npm/remixicon@4.1.0/fonts/remixicon.css" rel="stylesheet">
    <link href="https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700;800&display=swap" rel="stylesheet">
    
    <style>
    html, body, [class*="css"] { font-family: 'Nunito', sans-serif; color: #2D3748; }
    :root { --brand-blue: #004E92; --brand-coral: #FF6B6B; --card-radius: 16px; }
    
    div[data-baseweb="tab-highlight"] { background-color: transparent !important; }

    .unified-card {
        background-color: white; padding: 25px; border-radius: var(--card-radius);
        border: 1px solid #EDF2F7; box-shadow: 0 4px 6px rgba(0,0,0,0.03); margin-bottom: 20px;
    }
    
    .interactive-card:hover {
        transform: translateY(-3px); border-color: var(--brand-blue); box-shadow: 0 8px 15px rgba(0,78,146,0.08);
    }

    .header-clean {
        background-color: white; padding: 35px 40px; border-radius: var(--card-radius);
        border: 1px solid #EDF2F7; box-shadow: 0 4px 12px rgba(0,0,0,0.04); margin-bottom: 30px;
        display: flex; align-items: center; gap: 30px;
    }

    .stTabs [data-baseweb="tab-list"] { gap: 8px; padding-bottom: 10px; flex-wrap: wrap; }
    .stTabs [data-baseweb="tab"] {
        height: 40px; border-radius: 20px; padding: 0 20px; background-color: white;
        border: 1px solid #E2E8F0; font-weight: 700; color: #718096; font-size: 0.9rem;
    }
    .stTabs [aria-selected="true"] {
        background-color: var(--brand-coral) !important; color: white !important;
        border-color: var(--brand-coral) !important; box-shadow: 0 4px 10px rgba(255, 107, 107, 0.2);
    }

    .stTooltipIcon { color: var(--brand-blue) !important; cursor: help; }

    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] {
        border-radius: 12px !important; border-color: #E2E8F0 !important;
    }
    div[data-testid="column"] .stButton button {
        border-radius: 12px !important; font-weight: 800 !important; text-transform: uppercase; height: 50px !important; letter-spacing: 0.5px;
    }
    .icon-box {
        width: 48px; height: 48px; background: #EBF8FF; border-radius: 12px;
        display: flex; align-items: center; justify-content: center; margin-bottom: 15px;
        color: var(--brand-blue); font-size: 24px;
    }
    </style>
""", unsafe_allow_html=True)

# --- 4. IA (PROMPT PEDAGÓGICO SEM TÍTULO REDUNDANTE) ---
def consultar_gpt_v4(api_key, dados, contexto_pdf=""):
    if not api_key: return None, "⚠️ Configure a Chave API OpenAI na barra lateral."
    
    try:
        client = OpenAI(api_key=api_key)
        contexto_seguro = contexto_pdf[:5000] if contexto_pdf else "Sem laudo anexado."
        
        # Limpa interrogações
        evidencias_texto = "\n".join([f"- {k.replace('?', '')}" for k, v in dados['checklist_evidencias'].items() if v])
        
        meds_texto = ""
        if dados['lista_medicamentos']:
            for m in dados['lista_medicamentos']:
                meds_texto += f"- {m['nome']} ({m['posologia']}). Admin na escola: {'SIM' if m['escola'] else 'NÃO'}.\n"
        else: meds_texto = "Nenhuma medicação informada."

        mapeamento_texto = ""
        for categoria, itens in dados['barreiras_selecionadas'].items():
            if itens:
                mapeamento_texto += f"\n[{categoria}]: "
                detalhes = []
                for item in itens:
                    nivel = dados['niveis_suporte'].get(f"{categoria}_{item}", "Monitorado")
                    detalhes.append(f"{item} (Suporte {nivel})")
                mapeamento_texto += ", ".join(detalhes)

        prompt_sistema = """
        Você é um Neuropsicopedagogo Sênior.
        Sua missão é construir um PEI (Plano de Ensino Individualizado) centrado no estudante.
        
        REGRAS DE FORMATAÇÃO (CRUCIAL):
        1. NÃO COLOQUE TÍTULO NO DOCUMENTO (Ex: "PEI", "Plano de Ensino", "Relatório"). O cabeçalho já existe no PDF.
        2. Comece ESTRITAMENTE pelo título do tópico "1. PERFIL BIOPSICOSSOCIAL DO ESTUDANTE".
        3. Use CAIXA ALTA apenas nos títulos numéricos (1., 2., 3., 4., 5.).
        
        ESTRUTURA OBRIGATÓRIA:
        1. PERFIL BIOPSICOSSOCIAL DO ESTUDANTE (Narrativa humanizada)
        2. PLANEJAMENTO CURRICULAR E BNCC (Habilidades Essenciais + Recomposição)
        3. DIRETRIZES PRÁTICAS PARA ADAPTAÇÃO (Foco no Hiperfoco)
        4. PLANO DE INTERVENÇÃO E ESTRATÉGIAS
        5. PARECER FINAL E RECOMENDAÇÕES
        """

        prompt_usuario = f"""
        ESTUDANTE: {dados['nome']} | Série: {dados['serie']} | Turma: {dados['turma']}
        DIAGNÓSTICO: {dados['diagnostico']} (Se vazio, busque no laudo).
        MEDICAÇÕES: {meds_texto}
        
        QUEM É O ESTUDANTE:
        Histórico: {dados['historico']}
        Família: {dados['familia']}
        
        EVIDÊNCIAS DE SALA: {evidencias_texto}
        BARREIRAS: {mapeamento_texto}
        POTENCIALIDADES: Hiperfoco: {dados['hiperfoco']} | Fortes: {', '.join(dados['potencias'])}
        
        ESTRATÉGIAS: 
        Acesso: {', '.join(dados['estrategias_acesso'])}
        Ensino: {', '.join(dados['estrategias_ensino'])}
        Avaliação: {', '.join(dados['estrategias_avaliacao'])}
        
        LAUDO: {contexto_seguro}
        
        GERE O RELATÓRIO TÉCNICO. Comece direto pelo tópico 1.
        """
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": prompt_sistema}, {"role": "user", "content": prompt_usuario}],
            temperature=0.7
        )
        return response.choices[0].message.content, None
    except Exception as e: return None, f"Erro OpenAI: {str(e)}."

# --- 5. PDF (FORMATADOR INTELIGENTE) ---
class PDF_V3(FPDF):
    def header(self):
        self.set_draw_color(0, 78, 146); self.set_line_width(0.4)
        self.rect(5, 5, 200, 287)
        
        logo = finding_logo()
        if logo: 
            self.image(logo, 10, 10, 30)
            x_offset = 45 
        else: 
            x_offset = 12
        
        self.set_xy(x_offset, 16) 
        self.set_font('Arial', 'B', 16)
        self.set_text_color(0, 78, 146)
        self.cell(0, 8, 'PLANO DE ENSINO INDIVIDUALIZADO', 0, 1, 'L')
        
        self.set_xy(x_offset, 23)
        self.set_font('Arial', 'I', 10)
        self.set_text_color(100)
        self.cell(0, 5, 'Documento Oficial de Planejamento Pedagógico', 0, 1, 'L')
        self.ln(20)

    def footer(self):
        self.set_y(-15); self.set_font('Arial', 'I', 8); self.set_text_color(128)
        self.cell(0, 10, f'Gerado via PEI 360º | Página {self.page_no()}', 0, 0, 'C')

    def section_title(self, label):
        self.ln(8)
        self.set_fill_color(240, 248, 255)
        self.set_text_color(0, 78, 146)
        self.set_font('Arial', 'B', 11)
        self.cell(0, 8, f"  {label}", 0, 1, 'L', fill=True)
        self.ln(4)

def gerar_pdf(dados, tem_anexo):
    pdf = PDF_V3(); pdf.add_page(); pdf.set_auto_page_break(auto=True, margin=20)
    
    # 1. Identificação
    pdf.section_title("1. IDENTIFICAÇÃO E CONTEXTO")
    pdf.set_font("Arial", size=10); pdf.set_text_color(0)
    
    med_str = "; ".join([f"{m['nome']} ({m['posologia']})" for m in dados['lista_medicamentos']]) if dados['lista_medicamentos'] else "Não informado / Não faz uso."
    diag = dados['diagnostico'] if dados['diagnostico'] else ("Vide análise detalhada no parecer técnico" if tem_anexo else "Não informado")
    
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Nome:", 0, 0); pdf.set_font("Arial", '', 10); pdf.cell(0, 6, dados['nome'], 0, 1)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Nascimento:", 0, 0); pdf.set_font("Arial", '', 10); pdf.cell(0, 6, str(dados['nasc']), 0, 1)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Série/Turma:", 0, 0); pdf.set_font("Arial", '', 10); pdf.cell(0, 6, f"{dados['serie']} - {dados['turma']}", 0, 1)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Diagnóstico:", 0, 0); pdf.set_font("Arial", '', 10); pdf.multi_cell(0, 6, diag)
    pdf.ln(2)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Medicação:", 0, 0); pdf.set_font("Arial", '', 10); pdf.multi_cell(0, 6, med_str)
    pdf.ln(2)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Família:", 0, 0); pdf.set_font("Arial", '', 10); pdf.multi_cell(0, 6, dados['composicao_familiar'])

    # 2. Evidências
    evidencias = [k.replace('?', '') for k, v in dados['checklist_evidencias'].items() if v]
    if evidencias:
        pdf.section_title("2. PONTOS DE ATENÇÃO (EVIDÊNCIAS OBSERVADAS)")
        pdf.set_font("Arial", size=10)
        pdf.multi_cell(0, 6, limpar_texto_pdf('; '.join(evidencias) + '.'))

    # 3. Mapeamento
    tem_barreiras = any(dados['barreiras_selecionadas'].values())
    if tem_barreiras:
        pdf.section_title("3. MAPEAMENTO DE BARREIRAS E NÍVEIS DE SUPORTE")
        pdf.set_font("Arial", size=10)
        for categoria, itens in dados['barreiras_selecionadas'].items():
            if itens:
                pdf.set_font("Arial", 'B', 10)
                pdf.cell(0, 6, f"{categoria}:", 0, 1)
                pdf.set_font("Arial", size=10)
                for item in itens:
                    nivel = dados['niveis_suporte'].get(f"{categoria}_{item}", "Monitorado")
                    pdf.cell(5); pdf.cell(0, 6, f"- {item}: Suporte {nivel}", 0, 1)
                pdf.ln(2)

    # 4. Relatório IA
    if dados['ia_sugestao']:
        pdf.ln(5)
        # FORÇA COR E FONTE PARA EVITAR FORMATACAO ERRADA
        pdf.set_text_color(0) 
        pdf.set_font("Arial", '', 10)
        
        linhas = dados['ia_sugestao'].split('\n')
        for linha in linhas:
            linha_limpa = limpar_texto_pdf(linha)
            # Título Principal (1. TÍTULO)
            if re.match(r'^[1-5]\.', linha_limpa.strip()) and linha_limpa.strip().isupper():
                pdf.ln(4)
                pdf.set_fill_color(240, 248, 255)
                pdf.set_text_color(0, 78, 146)
                pdf.set_font('Arial', 'B', 11)
                pdf.cell(0, 8, f"  {linha_limpa}", 0, 1, 'L', fill=True)
                pdf.set_text_color(0) 
                pdf.set_font("Arial", size=10)
            # Subtítulos (Termina em :) ou linhas curtas
            elif linha_limpa.strip().endswith(':') and len(linha_limpa) < 70:
                pdf.ln(2)
                pdf.set_font("Arial", 'B', 10)
                pdf.multi_cell(0, 6, linha_limpa)
                pdf.set_font("Arial", size=10)
            # Texto comum
            else:
                pdf.multi_cell(0, 6, linha_limpa)
        
    pdf.ln(25)
    y = pdf.get_y()
    if y > 250: pdf.add_page(); y = 40
    pdf.line(20, y, 90, y); pdf.line(120, y, 190, y)
    pdf.set_font("Arial", 'I', 8); pdf.text(35, y+5, "Coordenação / Direção"); pdf.text(135, y+5, "Família / Responsável")
    return pdf.output(dest='S').encode('latin-1', 'replace')

def gerar_docx(dados):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(11)
    doc.add_heading('PLANO DE ENSINO INDIVIDUALIZADO', 0)
    doc.add_paragraph(f"Estudante: {dados['nome']} | Série: {dados['serie']}")
    if dados['ia_sugestao']: doc.add_heading('Parecer Técnico', level=1); doc.add_paragraph(dados['ia_sugestao'])
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0); return buffer

# --- 6. ESTADO E AUTO-REPARO ---
default_state = {
    'nome': '', 'nasc': date(2015, 1, 1), 'serie': None, 'turma': '', 'diagnostico': '', 
    'lista_medicamentos': [], 
    'composicao_familiar': '', 'historico': '', 'familia': '', 'hiperfoco': '', 'potencias': [],
    'rede_apoio': [], 'orientacoes_especialistas': '',
    'checklist_evidencias': {}, 
    'barreiras_selecionadas': {'Cognitivo': [], 'Comunicacional': [], 'Socioemocional': [], 'Sensorial/Motor': [], 'Acadêmico': []},
    'niveis_suporte': {}, 
    'estrategias_acesso': [], 'estrategias_ensino': [], 'estrategias_avaliacao': [], 'ia_sugestao': ''
}

if 'dados' not in st.session_state: st.session_state.dados = default_state
else:
    for key, val in default_state.items():
        if key not in st.session_state.dados: st.session_state.dados[key] = val

if 'pdf_text' not in st.session_state: st.session_state.pdf_text = ""

# --- 7. SIDEBAR ---
with st.sidebar:
    logo = finding_logo()
    if logo: st.image(logo, width=120)
    if 'OPENAI_API_KEY' in st.secrets: api_key = st.secrets['OPENAI_API_KEY']; st.success("✅ OpenAI OK")
    else: api_key = st.text_input("Chave OpenAI:", type="password")
    st.markdown("---")
    data_atual = date.today().strftime("%d/%m/%Y")
    st.markdown(f"<div style='font-size:0.75rem; color:#A0AEC0;'><b>PEI 360º</b><br>Criado e desenvolvido por<br><b>Rodrigo Amorim Queiroz</b><br>Atualizado: {data_atual}</div>", unsafe_allow_html=True)

# --- 8. LAYOUT ---
logo_path = finding_logo(); b64_logo = get_base64_image(logo_path); mime = "image/png"
img_html = f'<img src="data:{mime};base64,{b64_logo}" style="height: 80px;">' if logo_path else ""
st.markdown(f"""<div class="header-clean">{img_html}<div><p style="margin:0; color:#004E92; font-size:1.3rem; font-weight:800;">Ecossistema de Inteligência Pedagógica e Inclusiva</p></div></div>""", unsafe_allow_html=True)

abas = ["Início", "Estudante", "Coleta de Evidências", "Rede de Apoio", "Potencialidades & Barreiras", "Plano de Ação", "Consultoria IA", "Documento"]
tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(abas)

# TAB 0: INÍCIO
with tab0:
    st.markdown("### <i class='ri-dashboard-line'></i> Visão Geral", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1: st.markdown("""<div class="unified-card interactive-card"><div class="icon-box"><i class="ri-book-read-line"></i></div><h4>PEI 360º</h4><p>Sistema baseado em evidências para construção de Planos de Ensino Individualizados robustos.</p></div>""", unsafe_allow_html=True)
    with c2: st.markdown("""<div class="unified-card interactive-card"><div class="icon-box"><i class="ri-scales-3-line"></i></div><h4>Conformidade Legal</h4><p>Atende ao Decreto 12.686/2025: Foco nas barreiras e no nível de suporte, independente de laudo.</p></div>""", unsafe_allow_html=True)
    st.write("")
    c3, c4 = st.columns(2)
    with c3: st.markdown("""<div class="unified-card interactive-card"><div class="icon-box"><i class="ri-brain-line"></i></div><h4>Neurociência</h4><p>Mapeamos Funções Executivas e Perfil Sensorial para estratégias assertivas.</p></div>""", unsafe_allow_html=True)
    with c4: st.markdown("""<div class="unified-card interactive-card"><div class="icon-box"><i class="ri-compass-3-line"></i></div><h4>BNCC</h4><p>Garantia das Aprendizagens Essenciais através da flexibilização curricular.</p></div>""", unsafe_allow_html=True)

# TAB 1: ESTUDANTE
with tab1:
    st.markdown("### <i class='ri-user-star-line'></i> Dossiê do Estudante", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns([3, 2, 2, 1])
    st.session_state.dados['nome'] = c1.text_input("Nome Completo", st.session_state.dados['nome'])
    st.session_state.dados['nasc'] = c2.date_input("Nascimento", value=st.session_state.dados.get('nasc', date(2015, 1, 1)), min_value=date(2000, 1, 1), max_value=date.today())
    
    lista_series = ["Educação Infantil", "1º Ano (Anos Iniciais)", "2º Ano (Anos Iniciais)", "3º Ano (Anos Iniciais)", "4º Ano (Anos Iniciais)", "5º Ano (Anos Iniciais)", "6º Ano (Anos Finais)", "7º Ano (Anos Finais)", "8º Ano (Anos Finais)", "9º Ano (Anos Finais)", "1ª Série (Ensino Médio)", "2ª Série (Ensino Médio)", "3ª Série (Ensino Médio)"]
    st.session_state.dados['serie'] = c3.selectbox("Série/Ano", lista_series, placeholder="Selecione...")
    st.session_state.dados['turma'] = c4.text_input("Turma", st.session_state.dados['turma'])

    st.markdown("---")
    st.markdown("##### 1. Histórico e Contexto Familiar")
    ch1, ch2 = st.columns(2)
    with ch1:
        st.session_state.dados['historico'] = st.text_area("Histórico Escolar", st.session_state.dados['historico'], height=100, help="Trajetória escolar, retenções e avanços.")
    with ch2:
        st.session_state.dados['familia'] = st.text_area("Contexto Familiar", st.session_state.dados['familia'], height=100, help="Rotina, cuidadores e expectativas.")
    
    st.session_state.dados['composicao_familiar'] = st.text_input("Composição Familiar", st.session_state.dados.get('composicao_familiar', ''), placeholder="Quem mora com a criança?")

    st.markdown("##### 2. Saúde e Diagnóstico")
    st.session_state.dados['diagnostico'] = st.text_input("Diagnóstico Clínico (ou em investigação)", st.session_state.dados['diagnostico'], help="Se vazio, a IA tentará extrair do laudo anexo.")
    
    with st.container(border=True):
        st.markdown("**Controle de Medicação**")
        c_med1, c_med2, c_med3 = st.columns([3, 2, 1])
        with c_med1: novo_med = st.text_input("Nome do Medicamento", key="temp_med_nome")
        with c_med2: nova_pos = st.text_input("Posologia", key="temp_med_pos")
        with c_med3: 
            st.write("")
            st.write("")
            add_btn = st.button("➕ Adicionar")

        if add_btn and novo_med:
            st.session_state.dados['lista_medicamentos'].append({"nome": novo_med, "posologia": nova_pos, "escola": False})
            st.rerun()

        if st.session_state.dados['lista_medicamentos']:
            st.caption("Medicamentos Adicionados (Marque se administrar na escola):")
            for idx, med in enumerate(st.session_state.dados['lista_medicamentos']):
                c_list1, c_list2, c_list3 = st.columns([4, 2, 1])
                with c_list1: st.markdown(f"💊 **{med['nome']}** - {med['posologia']}")
                with c_list2: 
                    med['escola'] = st.checkbox("Na Escola?", value=med['escola'], key=f"check_med_{idx}")
                with c_list3: 
                    if st.button("🗑️", key=f"del_med_{idx}"):
                        st.session_state.dados['lista_medicamentos'].pop(idx)
                        st.rerun()

    with st.expander("📎 Anexar Laudo (PDF)"):
        up = st.file_uploader("Arquivo PDF", type="pdf")
        if up: st.session_state.pdf_text = ler_pdf(up); st.success("PDF Anexado!")

# TAB 2: COLETA DE EVIDÊNCIAS
with tab2:
    st.markdown("### <i class='ri-file-search-line'></i> Coleta de Evidências (Observação Dirigida)", unsafe_allow_html=True)
    st.info("O que você observa no dia a dia? Essas informações guiarão a IA para estratégias práticas.")
    
    questoes = {
        "Desafios no Currículo e Aprendizagem": ["O aluno não avança mesmo com atividades adaptadas?", "Os objetivos parecem distantes da realidade dele?", "Dificuldade em generalizar o aprendizado?", "Dificuldade com interpretação de texto?"],
        "Atenção e Processamento da Informação": ["Se perde durante a atividade?", "Dificuldade de assimilar novos conceitos?", "Esquece rapidamente o que foi ensinado (memória)?", "Demora para iniciar ou finalizar tarefas?"],
        "Comportamento e Autonomia": ["Precisa de explicação constante (1:1)?", "Recusa atividades mesmo quando adaptadas?", "Baixa tolerância à frustração?", "Dificuldade de organização dos materiais?"]
    }
    
    c_ev1, c_ev2, c_ev3 = st.columns(3)
    if 'checklist_evidencias' not in st.session_state.dados: st.session_state.dados['checklist_evidencias'] = {}

    with c_ev1:
        st.markdown("**Desafios no Currículo**")
        for q in questoes["Desafios no Currículo e Aprendizagem"]:
            st.session_state.dados['checklist_evidencias'][q] = st.checkbox(q, value=st.session_state.dados['checklist_evidencias'].get(q, False))
    with c_ev2:
        st.markdown("**Atenção e Processamento**")
        for q in questoes["Atenção e Processamento da Informação"]:
            st.session_state.dados['checklist_evidencias'][q] = st.checkbox(q, value=st.session_state.dados['checklist_evidencias'].get(q, False))
    with c_ev3:
        st.markdown("**Comportamento e Autonomia**")
        for q in questoes["Comportamento e Autonomia"]:
            st.session_state.dados['checklist_evidencias'][q] = st.checkbox(q, value=st.session_state.dados['checklist_evidencias'].get(q, False))

# TAB 3: REDE DE APOIO
with tab3:
    st.markdown("### <i class='ri-team-line'></i> Rede de Apoio", unsafe_allow_html=True)
    st.session_state.dados['rede_apoio'] = st.multiselect("Profissionais:", ["Psicólogo", "Fonoaudiólogo", "TO", "Neuropediatra", "Psicopedagogo", "Professor Particular"], placeholder="Selecione...")
    st.session_state.dados['orientacoes_especialistas'] = st.text_area("Orientações Técnicas (Resumo)", placeholder="Recomendações dos especialistas...", height=150)

# TAB 4: POTENCIALIDADES E BARREIRAS
with tab4:
    st.markdown("### <i class='ri-map-pin-user-line'></i> Mapeamento Integral", unsafe_allow_html=True)
    
    with st.container(border=True):
        st.markdown("#### <i class='ri-lightbulb-flash-line' style='color:#004E92'></i> Potencialidades e Hiperfoco", unsafe_allow_html=True)
        c_pot1, c_pot2 = st.columns(2)
        st.session_state.dados['hiperfoco'] = c_pot1.text_input("Hiperfoco", placeholder="Ex: Minecraft, Dinossauros...")
        potencias_opts = ["Memória Visual", "Raciocínio Lógico", "Criatividade", "Oralidade", "Artes", "Liderança", "Esportes/Motricidade", "Música", "Tecnologia", "Empatia/Social", "Observação de Detalhes"]
        st.session_state.dados['potencias'] = c_pot2.multiselect("Pontos Fortes", potencias_opts, placeholder="Selecione...")

    st.divider()
    
    with st.container(border=True):
        st.markdown("#### <i class='ri-barricade-line' style='color:#FF6B6B'></i> Barreiras e Nível de Suporte", unsafe_allow_html=True)
        
        categorias = {
            "Cognitivo": ["Atenção", "Memória de Trabalho", "Controle Inibitório", "Flexibilidade Cognitiva", "Planejamento", "Velocidade de Processamento", "Raciocínio Abstrato"],
            "Comunicacional": ["Linguagem Receptiva", "Linguagem Expressiva", "Pragmática (Uso Social)", "Articulação", "Comunicação Alternativa"],
            "Socioemocional": ["Regulação Emocional", "Tolerância à Frustração", "Interação Social", "Compreensão de Regras Sociais", "Rigidez de Pensamento"],
            "Sensorial/Motor": ["Coordenação Fina", "Coordenação Ampla", "Hipersensibilidade Auditiva", "Hipersensibilidade Tátil", "Propriocepção", "Visual"],
            "Acadêmico": ["Alfabetização (Leitura/Escrita)", "Interpretação de Texto", "Cálculo e Raciocínio Matemático", "Grafia/Legibilidade"]
        }

        cols = st.columns(3)
        idx = 0
        for cat_nome, itens in categorias.items():
            with cols[idx % 3]:
                with st.container():
                    st.markdown(f"**{cat_nome}**")
                    selecionados = st.multiselect(f"Barreiras:", itens, key=f"multi_{cat_nome}", placeholder="Selecione...", help="Selecione apenas o que impacta a aprendizagem.")
                    st.session_state.dados['barreiras_selecionadas'][cat_nome] = selecionados
                    
                    if selecionados:
                        st.caption("Intensidade do apoio:")
                        for item in selecionados:
                            val = st.select_slider(f"{item}", ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value="Monitorado", key=f"slider_{cat_nome}_{item}")
                            st.session_state.dados['niveis_suporte'][f"{cat_nome}_{item}"] = val
            idx += 1

# TAB 5: PLANO DE AÇÃO
with tab5:
    st.markdown("### <i class='ri-tools-line'></i> Plano de Ação Estratégico", unsafe_allow_html=True)
    
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        with st.container(border=True):
            st.markdown("#### 1. Acesso (DUA)")
            st.session_state.dados['estrategias_acesso'] = st.multiselect("Recursos:", ["Tempo Estendido", "Apoio à Leitura e Escrita", "Material Ampliado", "Sala Silenciosa", "Tecnologia", "Pausas", "Mobiliário Adaptado"], placeholder="Selecione...")
    with col_b:
        with st.container(border=True):
            st.markdown("#### 2. Ensino")
            st.session_state.dados['estrategias_ensino'] = st.multiselect("Metodologia:", ["Fragmentação de Tarefas", "Pistas Visuais", "Mapas Mentais", "Aprendizagem Baseada em Projetos", "Ensino Híbrido", "Modelagem", "Instrução Explícita"], placeholder="Selecione...")
    with col_c:
        with st.container(border=True):
            st.markdown("#### 3. Avaliação")
            st.session_state.dados['estrategias_avaliacao'] = st.multiselect("Formato:", ["Prova Adaptada (Conteúdo)", "Consulta Permitida", "Avaliação Oral", "Sem Distratores Visuais", "Portfólio", "Autoavaliação"], placeholder="Selecione...")

# TAB 6: IA
with tab6:
    st.markdown("### <i class='ri-brain-line'></i> Consultoria Pedagógica", unsafe_allow_html=True)
    c1, c2 = st.columns([1, 2])
    with c1:
        st.markdown("""<div style="background:#F0F4FF; padding:15px; border-radius:12px; border-left: 4px solid #004E92; color: #2D3748; font-size: 0.95rem;">Olá! Sou sua assistente de inteligência pedagógica. Estou pronta para cruzar os dados do estudante com a <b>BNCC</b> e a <b>Neurociência</b>.</div>""", unsafe_allow_html=True)
        st.write("")
        with st.expander("🔍 Ver detalhes do processamento"):
            st.markdown("- **Análise de Evidências:** Leitura dos checklists.\n- **Matriz de Suporte:** Cruzamento das barreiras com o nível de ajuda.\n- **Legislação:** Validação com a LBI.")

        if st.button("GERAR PLANO AGORA", type="primary"):
            if not st.session_state.dados['nome']: st.error("Preencha o Nome do aluno.")
            else:
                with st.spinner("Processando dados pedagógicos..."):
                    res, err = consultar_gpt_v4(api_key, st.session_state.dados, st.session_state.pdf_text)
                    if err: st.error(err)
                    else: st.session_state.dados['ia_sugestao'] = res; st.success("Plano Gerado!")
    with c2:
        if st.session_state.dados['ia_sugestao']: st.text_area("Texto do Relatório (Editável):", st.session_state.dados['ia_sugestao'], height=600)

# TAB 7: DOCUMENTO
with tab7:
    st.markdown("### <i class='ri-file-pdf-line'></i> Exportação Oficial", unsafe_allow_html=True)
    if st.session_state.dados['ia_sugestao']:
        c1, c2 = st.columns(2)
        tem_anexo = len(st.session_state.pdf_text) > 0
        with c1:
            pdf = gerar_pdf(st.session_state.dados, tem_anexo)
            st.download_button("📥 Baixar PDF", pdf, f"PEI_{st.session_state.dados['nome']}.pdf", "application/pdf", type="primary")
        with c2:
            docx = gerar_docx(st.session_state.dados)
            st.download_button("📥 Baixar Word", docx, f"PEI_{st.session_state.dados['nome']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else: st.warning("Gere o plano na aba Consultoria IA primeiro.")

st.markdown("---")
st.markdown("<div style='text-align: center; color: #A0AEC0; font-size: 0.8rem;'>PEI 360º v5.2 | Perfect Print Edition</div>", unsafe_allow_html=True)