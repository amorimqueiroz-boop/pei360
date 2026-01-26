import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openai import OpenAI
from pypdf import PdfReader
from fpdf import FPDF
import base64
import os
import re
import hashlib

# ==================== SISTEMA DE AUTENTICAÇÃO ====================
AUTH_PASSWORD_HASH = "8c6976e5b5410415bde908bd4dee15dfb167a9c873fc4bb8a81f6f2ab448a918"  # hash SHA-256 de "admin"
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.markdown("""
    <style>
    .login-container {
        max-width: 400px;
        margin: 100px auto;
        padding: 40px;
        background: white;
        border-radius: 20px;
        box-shadow: 0 10px 40px rgba(0,0,0,0.1);
    }
    .login-title {
        color: #004E92;
        text-align: center;
        margin-bottom: 30px;
        font-weight: 800;
    }
    </style>
    """, unsafe_allow_html=True)
    
    with st.container():
        st.markdown('<div class="login-container">', unsafe_allow_html=True)
        st.markdown('<h2 class="login-title">🔐 Acesso Restrito</h2>', unsafe_allow_html=True)
        password = st.text_input("Senha de acesso:", type="password", key="password_input")
        
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("Entrar", use_container_width=True, type="primary"):
                if hashlib.sha256(password.encode()).hexdigest() == AUTH_PASSWORD_HASH:
                    st.session_state.authenticated = True
                    st.rerun()
                else:
                    st.error("Senha incorreta")
        with col2:
            if st.button("Sair", use_container_width=True):
                st.stop()
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("""
        <div style='text-align:center; margin-top:30px; color:#718096; font-size:0.8rem;'>
        Sistema protegido por autenticação<br>
        Contate o administrador para acesso
        </div>
        """, unsafe_allow_html=True)
        st.stop()

# ==================== CÓDIGO PRINCIPAL (OFUSCADO) ====================
def a1():
    if os.path.exists("iconeaba.png"): return "iconeaba.png"
    return "📘"

st.set_page_config(page_title="PEI 360º", page_icon=a1(), layout="wide", initial_sidebar_state="expanded")

def a2():
    p = ["360.png", "360.jpg", "logo.png", "logo.jpg", "iconeaba.png"]
    for n in p:
        if os.path.exists(n): return n
    return None

def a3(i):
    if not i: return ""
    with open(i, "rb") as f: return base64.b64encode(f.read()).decode()

def a4(a):
    if a is None: return ""
    try:
        r = PdfReader(a)
        t = ""
        for i, p in enumerate(r.pages):
            if i >= 6: break
            t += p.extract_text() + "\n"
        return t
    except: return ""

def a5(t):
    if not t: return ""
    t = t.replace('**', '').replace('__', '')
    t = t.replace('### ', '').replace('## ', '').replace('# ', '')
    t = t.replace('* ', '• ')
    t = re.sub(r'[^\x00-\xff]', '', t)
    return t

st.markdown("""<link href="https://cdn.jsdelivr.net/npm/remixicon@4.1.0/fonts/remixicon.css" rel="stylesheet"><link href="https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700;800&display=swap" rel="stylesheet"><style>html,body,[class*="css"]{font-family:'Nunito',sans-serif;color:#2D3748;}:root{--brand-blue:#004E92;--brand-coral:#FF6B6B;--bg-gray:#F7FAFC;--card-radius:16px;}div[data-baseweb="tab-highlight"]{background-color:transparent!important;}.unified-card{background-color:white;padding:25px;border-radius:var(--card-radius);border:1px solid #EDF2F7;box-shadow:0 4px 6px rgba(0,0,0,0.03);margin-bottom:20px;}.interactive-card:hover{transform:translateY(-3px);border-color:var(--brand-blue);box-shadow:0 8px 15px rgba(0,78,146,0.08);}.header-content{display:flex;align-items:center;gap:25px;}.stTabs [data-baseweb="tab-list"]{gap:10px;padding-bottom:10px;}.stTabs [data-baseweb="tab"]{height:45px;border-radius:25px;padding:0 25px;background-color:white;border:1px solid #E2E8F0;font-weight:700;color:#718096;}.stTabs [aria-selected="true"]{background-color:var(--brand-coral)!important;color:white!important;border-color:var(--brand-coral)!important;box-shadow:0 4px 10px rgba(255,107,107,0.2);}.icon-box{width:45px;height:45px;background:#EBF8FF;border-radius:12px;display:flex;align-items:center;justify-content:center;margin-bottom:15px;color:var(--brand-blue);font-size:22px;}.stTextInput input,.stTextArea textarea,.stSelectbox div[data-baseweb="select"]{border-radius:12px!important;border-color:#E2E8F0!important;}div[data-testid="column"] .stButton button{border-radius:12px!important;font-weight:800!important;text-transform:uppercase;height:50px!important;letter-spacing:0.5px;}</style>""", unsafe_allow_html=True)

def a6(k, d, c=""):
    if not k: return None, "⚠️ Configure a Chave API."
    try:
        client = OpenAI(api_key=k)
        cs = c[:5000] if c else "Sem laudo."
        is_ahsd = "altas habilidades" in d['diagnostico'].lower() or "superdotação" in d['diagnostico'].lower()
        f = "ENRIQUECIMENTO" if is_ahsd else "FLEXIBILIZAÇÃO"
        ps = "Neuropsicopedagogo Sênior. Tarefa: Redigir o PEI."
        pu = f"""
        ESTUDANTE: {d['nome']} | Série: {d['serie']}
        DIAGNÓSTICO: {d['diagnostico']} ({f})
        MEDICAÇÃO: {d['medicacao']}
        HIPERFOCO: {d['hiperfoco']}
        PONTOS FORTES: {', '.join(d['potencias'])}
        CONTEXTO: {d['historico']} | {d['familia']}
        REDE: {', '.join(d['rede_apoio'])} | {d['orientacoes_especialistas']}
        BARREIRAS SENSORIAL: {', '.join(d['b_sensorial'])}
        BARREIRAS COGNITIVO: {', '.join(d['b_cognitiva'])}
        BARREIRAS SOCIAL: {', '.join(d['b_social'])}
        ESTRATÉGIAS ACESSO: {', '.join(d['estrategias_acesso'])}
        ESTRATÉGIAS ENSINO: {', '.join(d['estrategias_ensino'])}
        ESTRATÉGIAS AVALIAÇÃO: {', '.join(d['estrategias_avaliacao'])}
        LAUDO PDF: {cs}
        """
        r = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "system", "content": ps}, {"role": "user", "content": pu}], temperature=0.7)
        return r.choices[0].message.content, None
    except Exception as e: return None, f"Erro: {str(e)}."

class PDF(FPDF):
    def header(self):
        self.set_draw_color(0, 78, 146); self.set_line_width(0.4); self.rect(5, 5, 200, 287)
        logo = a2()
        if logo: self.image(logo, 12, 12, 22); x = 40
        else: x = 12
        self.set_xy(x, 15); self.set_font('Arial', 'B', 14); self.set_text_color(0, 78, 146); self.cell(0, 8, 'PLANO DE ENSINO INDIVIDUALIZADO', 0, 1, 'L')
        self.set_xy(x, 22); self.set_font('Arial', 'I', 9); self.set_text_color(100); self.cell(0, 5, 'Documento Oficial', 0, 1, 'L'); self.ln(15)
    def footer(self):
        self.set_y(-15); self.set_font('Arial', 'I', 8); self.set_text_color(128); self.cell(0, 10, f'Página {self.page_no()}', 0, 0, 'C')
    def section_title(self, l):
        self.ln(5); self.set_fill_color(240, 248, 255); self.set_text_color(0, 78, 146); self.set_font('Arial', 'B', 11); self.cell(0, 8, f"  {l}", 0, 1, 'L', fill=True); self.ln(3)

def a7(d, a):
    pdf = PDF(); pdf.add_page(); pdf.set_auto_page_break(auto=True, margin=20)
    pdf.section_title("1. IDENTIFICAÇÃO")
    pdf.set_font("Arial", size=10); pdf.set_text_color(0)
    n = d['nasc'].strftime('%d/%m/%Y') if d['nasc'] else "-"
    dd = d['diagnostico'] if d['diagnostico'] else ("Em análise" if a else "Não informado")
    md = d['medicacao'] if d['medicacao'] else "Não faz uso"
    t = f"Nome: {d['nome']}\nNascimento: {n}\nSérie: {d['serie']} | Turma: {d['turma']}\nDiagnóstico: {dd}\nMedicação: {md}"
    pdf.multi_cell(0, 6, a5(t))
    if d['rede_apoio'] or d['orientacoes_especialistas']:
        pdf.ln(3); pdf.set_font("Arial", 'B', 10); pdf.cell(0, 6, "Suporte:", 0, 1); pdf.set_font("Arial", size=10)
        p = ', '.join(d['rede_apoio']) if d['rede_apoio'] else "-"; o = d['orientacoes_especialistas'] if d['orientacoes_especialistas'] else "-"
        pdf.multi_cell(0, 6, a5(f"Profissionais: {p}.\nOrientações: {o}"))
    if d['ia_sugestao']: pdf.ln(5); pdf.multi_cell(0, 6, a5(d['ia_sugestao']))
    pdf.ln(20); y = pdf.get_y(); if y > 250: pdf.add_page(); y = 40
    pdf.line(20, y, 90, y); pdf.line(120, y, 190, y); pdf.set_font("Arial", 'I', 8); pdf.text(35, y+5, "Coordenação"); pdf.text(135, y+5, "Família")
    return pdf.output(dest='S').encode('latin-1', 'replace')

def a8(d):
    doc = Document(); s = doc.styles['Normal']; s.font.name = 'Arial'; s.font.size = Pt(11)
    doc.add_heading('PLANO DE ENSINO INDIVIDUALIZADO', 0)
    doc.add_paragraph(f"Estudante: {d['nome']}"); doc.add_paragraph(f"Série: {d['serie']} | Turma: {d['turma']}")
    doc.add_paragraph(f"Diagnóstico: {d['diagnostico']}"); doc.add_paragraph(f"Medicação: {d['medicacao']}")
    if d['ia_sugestao']: doc.add_heading('Parecer Pedagógico', level=1); doc.add_paragraph(d['ia_sugestao'])
    b = BytesIO(); doc.save(b); b.seek(0); return b

if 'dados' not in st.session_state:
    st.session_state.dados = {'nome': '', 'nasc': None, 'serie': None, 'turma': '', 'diagnostico': '', 'medicacao': '', 'historico': '', 'familia': '', 'hiperfoco': '', 'potencias': [], 'rede_apoio': [], 'orientacoes_especialistas': '', 'b_sensorial': [], 'sup_sensorial': '🟡 Monitorado', 'b_cognitiva': [], 'sup_cognitiva': '🟡 Monitorado', 'b_social': [], 'sup_social': '🟡 Monitorado', 'estrategias_acesso': [], 'estrategias_ensino': [], 'estrategias_avaliacao': [], 'ia_sugestao': ''}
if 'pdf_text' not in st.session_state: st.session_state.pdf_text = ""

with st.sidebar:
    logo = a2()
    if logo: st.image(logo, width=120)
    if 'OPENAI_API_KEY' in st.secrets: api_key = st.secrets['OPENAI_API_KEY']; st.success("✅ OpenAI Ativa")
    else: api_key = st.text_input("Chave OpenAI:", type="password")
    st.markdown("---"); st.markdown("<div style='font-size:0.8rem; color:#A0AEC0;'>PEI 360º v3.5</div>", unsafe_allow_html=True)

logo_path = a2(); b64_logo = a3(logo_path); mime = "image/png" if logo_path and logo_path.endswith("png") else "image/jpeg"
img_html = f'<img src="data:{mime};base64,{b64_logo}" style="height: 70px;">' if logo_path else ""
st.markdown(f'<div class="unified-card header-content">{img_html}<div><p style="margin: 0; color: #004E92; font-size: 1.2rem; font-weight: 700;">Ecossistema de Inteligência Pedagógica</p></div></div>', unsafe_allow_html=True)

abas = ["Início", "Estudante", "Rede de Apoio", "Mapeamento", "Plano de Ação", "Consultoria IA", "Documento"]
tab0, tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(abas)

with tab0:
    st.markdown("### <i class='ri-dashboard-line'></i> Visão Geral", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1: st.markdown('<div class="unified-card interactive-card"><div class="icon-box"><i class="ri-book-read-line"></i></div><h4>O que é o PEI?</h4><p>Instrumento oficial de acessibilidade curricular.</p></div>', unsafe_allow_html=True)
    with c2: st.markdown('<div class="unified-card interactive-card"><div class="icon-box"><i class="ri-scales-3-line"></i></div><h4>Legalidade</h4><p>Conforme Decreto 12.686/2025.</p></div>', unsafe_allow_html=True)
    c3, c4 = st.columns(2)
    with c3: st.markdown('<div class="unified-card interactive-card"><div class="icon-box"><i class="ri-brain-line"></i></div><h4>Neurociência</h4><p>Mapeamos Funções Executivas.</p></div>', unsafe_allow_html=True)
    with c4: st.markdown('<div class="unified-card interactive-card"><div class="icon-box"><i class="ri-compass-3-line"></i></div><h4>BNCC</h4><p>Garantia das Aprendizagens Essenciais.</p></div>', unsafe_allow_html=True)

with tab1:
    st.markdown("### <i class='ri-user-smile-line'></i> Dossiê do Estudante", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns([3, 2, 2, 1])
    st.session_state.dados['nome'] = c1.text_input("Nome Completo", st.session_state.dados['nome'])
    st.session_state.dados['nasc'] = c2.date_input("Nascimento", st.session_state.dados['nasc'])
    st.session_state.dados['serie'] = c3.selectbox("Série/Ano", ["Infantil", "1º Ano", "2º Ano", "3º Ano", "4º Ano", "5º Ano", "Fund. II", "Ensino Médio"])
    st.session_state.dados['turma'] = c4.text_input("Turma", st.session_state.dados['turma'])
    st.markdown("---"); st.markdown("##### 1. Contexto Escolar e Familiar")
    ch, cf = st.columns(2)
    with ch: st.session_state.dados['historico'] = st.text_area("Histórico Escolar", st.session_state.dados['historico'], height=100, label_visibility="collapsed")
    with cf: st.session_state.dados['familia'] = st.text_area("Contexto Familiar", st.session_state.dados['familia'], height=100, label_visibility="collapsed")
    st.markdown("##### 2. Saúde e Diagnóstico")
    col_d, col_m = st.columns(2)
    with col_d: st.session_state.dados['diagnostico'] = st.text_input("Diagnóstico Clínico", st.session_state.dados['diagnostico'])
    with col_m: st.session_state.dados['medicacao'] = st.text_input("Medicação em uso", st.session_state.dados['medicacao'])
    with st.expander("📎 Anexar Laudo (PDF)"):
        up = st.file_uploader("Arquivo PDF", type="pdf", key="pdf_uploader")
        if up: st.session_state.pdf_text = a4(up); st.success("PDF Anexado!")

with tab2:
    st.markdown("### <i class='ri-team-line'></i> Rede de Apoio", unsafe_allow_html=True)
    c_rede1, c_rede2 = st.columns(2)
    st.session_state.dados['rede_apoio'] = c_rede1.multiselect("Profissionais:", ["Psicólogo", "Fonoaudiólogo", "Terapeuta Ocupacional", "Neuropediatra", "Psicopedagogo", "Professor Particular"])
    st.session_state.dados['orientacoes_especialistas'] = st.text_area("Orientações Técnicas", placeholder="Recomendações...", height=150)

with tab3:
    st.markdown("### <i class='ri-map-pin-user-line'></i> Mapeamento Integral", unsafe_allow_html=True)
    with st.container(border=True):
        st.markdown("#### Potencialidades e Interesses")
        c_pot1, c_pot2 = st.columns(2)
        with c_pot1: st.session_state.dados['hiperfoco'] = st.text_input("Hiperfoco", placeholder="Ex: Dinossauros...")
        with c_pot2: st.session_state.dados['potencias'] = st.multiselect("Pontos Fortes", ["Memória Visual", "Lógica Matemática", "Criatividade", "Oralidade", "Tecnologia", "Artes", "Música"])
    st.markdown("#### Barreiras e Suporte")
    c_bar1, c_bar2, c_bar3 = st.columns(3)
    with c_bar1:
        with st.container(border=True):
            st.markdown("##### Sensorial")
            st.session_state.dados['b_sensorial'] = st.multiselect("Barreiras:", ["Hipersensibilidade Auditiva", "Hipersensibilidade Visual", "Busca Sensorial", "Baixo Tônus"], key="b1")
            st.session_state.dados['sup_sensorial'] = st.select_slider("Suporte", ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value="Monitorado", key="s1")
    with c_bar2:
        with st.container(border=True):
            st.markdown("##### Cognitivo")
            st.session_state.dados['b_cognitiva'] = st.multiselect("Barreiras:", ["Atenção", "Memória", "Rigidez Mental", "Processamento Lento"], key="b2")
            st.session_state.dados['sup_cognitiva'] = st.select_slider("Suporte", ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value="Monitorado", key="s2")
    with c_bar3:
        with st.container(border=True):
            st.markdown("##### Social")
            st.session_state.dados['b_social'] = st.multiselect("Barreiras:", ["Interação", "Frustração", "Regras", "Isolamento"], key="b3")
            st.session_state.dados['sup_social'] = st.select_slider("Suporte", ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value="Monitorado", key="s3")

with tab4:
    st.markdown("### <i class='ri-tools-line'></i> Estratégias Pedagógicas", unsafe_allow_html=True)
    c_acesso, c_ensino = st.columns(2)
    with c_acesso:
        st.markdown("#### 1. Acesso ao Currículo")
        st.session_state.dados['estrategias_acesso'] = st.multiselect("Recursos de Acessibilidade:", ["Tempo Estendido (+25%)", "Apoio à Leitura e Escrita", "Material Ampliado", "Sala com Redução de Estímulos", "Tecnologia Assistiva", "Pausas Sensoriais"])
    with c_ensino:
        st.markdown("#### 2. Metodologia de Ensino")
        st.session_state.dados['estrategias_ensino'] = st.multiselect("Estratégias Didáticas:", ["Fragmentação de Tarefas", "Pistas Visuais", "Enriquecimento Curricular", "Antecipação de Rotina", "Projetos Práticos"])
    st.write(""); st.markdown("#### 3. Avaliação")
    st.session_state.dados['estrategias_avaliacao'] = st.multiselect("Formato Avaliativo:", ["Prova Adaptada", "Consulta Permitida", "Avaliação Oral", "Trabalho Prático", "Enunciados Curtos"])

with tab5:
    st.markdown("### <i class='ri-robot-2-line'></i> Consultoria Pedagógica", unsafe_allow_html=True)
    col_btn, col_txt = st.columns([1, 2])
    with col_btn:
        if st.button("GERAR PLANO", type="primary"):
            if not st.session_state.dados['nome']: st.error("Preencha o Nome.")
            else:
                with st.spinner("Processando..."):
                    res, err = a6(api_key, st.session_state.dados, st.session_state.pdf_text)
                    if err: st.error(err)
                    else: st.session_state.dados['ia_sugestao'] = res; st.success("Gerado!")
    with col_txt:
        if st.session_state.dados['ia_sugestao']: st.text_area("Parecer Técnico:", st.session_state.dados['ia_sugestao'], height=500)
        else: st.markdown('<div style="padding:50px; text-align:center; color:#CBD5E0; border:2px dashed #E2E8F0; border-radius:12px;">O plano aparecerá aqui.</div>', unsafe_allow_html=True)

with tab6:
    st.markdown("### <i class='ri-file-pdf-line'></i> Exportação", unsafe_allow_html=True)
    if st.session_state.dados['ia_sugestao']:
        c_pdf, c_word = st.columns(2); a = len(st.session_state.pdf_text) > 0
        with c_pdf: pdf_bytes = a7(st.session_state.dados, a); st.download_button("📥 Baixar PDF", pdf_bytes, f"PEI_{st.session_state.dados['nome']}.pdf", "application/pdf", type="primary")
        with c_word: docx_bytes = a8(st.session_state.dados); st.download_button("📥 Baixar Word", docx_bytes, f"PEI_{st.session_state.dados['nome']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else: st.warning("Gere o plano na aba de IA primeiro.")

st.markdown("---"); st.markdown('<div style="text-align: center; color: #A0AEC0; font-size: 0.8rem;">PEI 360º v3.5</div>', unsafe_allow_html=True)
